"""Convert DEPLOYMENT.md into a formatted Word (.docx) document."""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SRC = Path(__file__).parent / "DEPLOYMENT.md"
OUT = Path(__file__).parent / "DEPLOYMENT.docx"

CODE_FONT = "Consolas"
CODE_SHADE = "F2F2F2"


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]

    shade = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{CODE_SHADE}"/>'
    from docx.oxml import parse_xml

    cell._tc.get_or_add_tcPr().append(parse_xml(shade))

    first = True
    for line in lines:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line if line else "")
        run.font.name = CODE_FONT
        run.font.size = Pt(9)


def add_formatted_runs(paragraph, text: str) -> None:
    """Handle inline bold (**text**) segments."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        hdr_cells[i].paragraphs[0].add_run(text.strip()).bold = True

    for row in body:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if i < len(cells):
                cells[i].text = text.strip()

    doc.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    return [c for c in line.strip().strip("|").split("|")]


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # Table
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            table_rows = [parse_table_row(line)]
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        # Bullet list
        elif re.match(r"^\s*-\s+", line):
            indent = len(line) - len(line.lstrip())
            content = re.sub(r"^\s*-\s+", "", line)
            style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
            para = doc.add_paragraph(style=style_name)
            add_formatted_runs(para, content)
        # Numbered list
        elif re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            add_formatted_runs(para, content)
        # Blank line
        elif line.strip() == "":
            pass
        # Normal paragraph
        else:
            para = doc.add_paragraph()
            add_formatted_runs(para, line)

        i += 1

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
